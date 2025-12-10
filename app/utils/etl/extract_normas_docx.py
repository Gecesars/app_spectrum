"""
Extrai tabelas normativas dos DOCX e gera CSVs em data/normas/.

Mapeamentos:
- FM (doc FM):
  * Tabela classes (idx=2) -> normas_fm_classes.csv
  * Tabela C/I (idx=4) -> normas_fm_protecao.csv
  * Tabela distâncias por classe (idx=5) -> normas_fm_radcom_distancias.csv (assumidas em km)
- TV (doc TV):
  * Níveis de contorno (idx=0) -> normas_tv_nivel_contorno.csv
  * Classes digitais (idx=1) -> normas_tv_digital_classes.csv
  * Classes analógicas (idx=2) -> normas_tv_analogica_classes.csv
  * C/I (idx=3) -> normas_tv_protecao.csv
"""

import csv
import os
from typing import List, Optional

import docx

BASE_DIR = os.path.join("data", "normas")

FM_DOC = os.path.join(
    BASE_DIR, "REQUISITOS TÉCNICOS DE CONDIÇÕES DE USO DE RADIOFREQUÊNCIAS PARA OS.docx"
)
TV_DOC = os.path.join(
    BASE_DIR,
    "REQUISITOS TÉCNICOS de cONDIÇÕES DE USO DE RADIOFREQUÊNCIAS PARA OS SERVIÇOS DE RADIODIFUSÃO DE SONS E IMAGENS E DE RETRANSMISSÃO DE TELEVISÃO.docx",
)


def to_float(value: str) -> Optional[float]:
    if value is None:
        return None
    v = value.strip().replace("%", "")
    if v in {"", "--", "não aplicável", "nao aplicavel", "não aplicavel"}:
        return None
    v = v.replace(",", ".")
    try:
        return float(v)
    except Exception:
        return None


def write_csv(path: str, header: List[str], rows: List[List]) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(header)
        for row in rows:
            writer.writerow(row)
    print(f"gravado {path} ({len(rows)} linhas)")


def extract_fm_classes(doc: docx.Document) -> None:
    tbl = doc.tables[2]
    rows = []
    for r in tbl.rows[3:]:
        cells = [c.text.strip() for c in r.cells]
        classe = cells[0]
        erp_kw = to_float(cells[1])
        dist_km = to_float(cells[3])
        hnmt = to_float(cells[4])
        if not classe:
            continue
        rows.append([classe, erp_kw, hnmt, dist_km])
    write_csv(
        os.path.join(BASE_DIR, "normas_fm_classes.csv"),
        ["classe", "erp_max_kw", "hnmt_max_m", "dist_max_contorno66_km"],
        rows,
    )


def extract_fm_protecao(doc: docx.Document) -> None:
    tbl = doc.tables[4]
    rows = []
    for r in tbl.rows[2:]:
        cells = [c.text.strip() for c in r.cells]
        tipo = "cocanal" if "COCANAL" in cells[0].upper() else "adjacente_200khz"
        delta = to_float(cells[1])
        ci_db = to_float(cells[3])
        rows.append([tipo, delta, ci_db])
    write_csv(
        os.path.join(BASE_DIR, "normas_fm_protecao.csv"),
        ["tipo_interferencia", "delta_f_khz", "ci_requerida_db"],
        rows,
    )


def extract_fm_radcom(doc: docx.Document) -> None:
    tbl = doc.tables[5]
    rows = []
    for r in tbl.rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        classe = cells[0]
        cocanal = to_float(cells[1])
        adj = to_float(cells[2])
        if classe:
            rows.append([classe, "cocanal", cocanal])
            rows.append([classe, "adjacente_1", adj])
    write_csv(
        os.path.join(BASE_DIR, "normas_fm_radcom_distancias.csv"),
        ["classe_fm", "situacao", "dist_min_km"],
        rows,
    )


def extract_tv_niveis(doc: docx.Document) -> None:
    tbl = doc.tables[0]
    rows = []
    header = ["tecnologia", "faixa_canal", "nivel_campo_dbuv_m"]

    faixa_map = {"2 a 6": "vhf_baixo", "7 a 13": "vhf_alto", "14 a 51": "uhf"}
    for r in tbl.rows[2:]:
        cells = [c.text.strip() for c in r.cells]
        tecnologia = "analogica" if "analóg" in cells[0].lower() else "digital"
        for idx in range(1, len(cells)):
            faixa = faixa_map.get(tbl.rows[1].cells[idx].text.strip(), None)
            if not faixa:
                continue
            val = to_float(cells[idx])
            if val is None:
                continue
            rows.append([tecnologia, faixa, val])
    write_csv(os.path.join(BASE_DIR, "normas_tv_nivel_contorno.csv"), header, rows)


def _explode_channels(chan_cell: str, val_cell: str) -> List[tuple]:
    chans = [c.strip() for c in chan_cell.split("\n") if c.strip()]
    vals = [v.strip() for v in val_cell.split("\n") if v.strip()]
    return list(zip(chans, vals))


def extract_tv_classes(doc: docx.Document) -> None:
    # Digital: table 1
    rows_dig = []
    tbl = doc.tables[1]
    for r in tbl.rows[1:]:
        classe = r.cells[0].text.strip()
        chans_cell = r.cells[1].text
        erp_cell = r.cells[2].text
        dist_cell = r.cells[4].text
        for (chan_range, erp_str), (_, dist_str) in zip(
            _explode_channels(chans_cell, erp_cell), _explode_channels(chans_cell, dist_cell)
        ):
            rows_dig.append(
                [
                    classe,
                    chan_range.replace("–", "-").replace("—", "-").replace(" ", ""),
                    to_float(erp_str.replace("kW", "")),
                    to_float(r.cells[3].text),
                    to_float(dist_str),
                ]
            )
    write_csv(
        os.path.join(BASE_DIR, "normas_tv_digital_classes.csv"),
        ["classe", "faixa_canal", "erp_max_kw", "hnmt_ref_m", "dist_max_contorno_protegido_km"],
        rows_dig,
    )

    # Analógica: table 2
    rows_an = []
    tbl = doc.tables[2]
    for r in tbl.rows[1:]:
        classe = r.cells[0].text.strip()
        chans_cell = r.cells[1].text
        erp_cell = r.cells[2].text
        dist_cell = r.cells[4].text
        for (chan_range, erp_str), (_, dist_str) in zip(
            _explode_channels(chans_cell, erp_cell), _explode_channels(chans_cell, dist_cell)
        ):
            rows_an.append(
                [
                    classe,
                    chan_range.replace("–", "-").replace("—", "-").replace(" ", ""),
                    to_float(erp_str.replace("kW", "")),
                    to_float(r.cells[3].text),
                    to_float(dist_str),
                ]
            )
    write_csv(
        os.path.join(BASE_DIR, "normas_tv_analogica_classes.csv"),
        ["classe", "faixa_canal", "erp_max_kw", "hnmt_ref_m", "dist_max_contorno_protegido_km"],
        rows_an,
    )


def extract_tv_protecao(doc: docx.Document) -> None:
    tbl = doc.tables[3]
    header_cols = tbl.rows[1].cells
    combos = [
        ("analógico", "analógico", header_cols[2].text),
        ("analógico", "digital", header_cols[3].text),
        ("digital", "analógico", header_cols[4].text),
        ("digital", "digital", header_cols[5].text),
    ]
    rows = []
    for r in tbl.rows[2:]:
        tipo = r.cells[0].text.strip()
        delta = r.cells[1].text.strip()
        for idx, (tec_d, tec_i, _) in enumerate(combos):
            val = to_float(r.cells[2 + idx].text)
            if val is None:
                continue
            rows.append(
                {
                    "tipo_interferencia": tipo.lower().replace(" ", "_"),
                    "tecnologia_desejado": "analogica" if "analóg" in tec_d else "digital",
                    "tecnologia_interferente": "analogica" if "analóg" in tec_i else "digital",
                    "delta_canal": delta,
                    "ci_requerida_db": val,
                    "observacao": None,
                }
            )
    write_csv(
        os.path.join(BASE_DIR, "normas_tv_protecao.csv"),
        [
            "tipo_interferencia",
            "tecnologia_desejado",
            "tecnologia_interferente",
            "delta_canal",
            "ci_requerida_db",
            "observacao",
        ],
        [[r[k] for k in ["tipo_interferencia", "tecnologia_desejado", "tecnologia_interferente", "delta_canal", "ci_requerida_db", "observacao"]] for r in rows],
    )


def run() -> None:
    fm_doc = docx.Document(FM_DOC)
    tv_doc = docx.Document(TV_DOC)

    extract_fm_classes(fm_doc)
    extract_fm_protecao(fm_doc)
    extract_fm_radcom(fm_doc)

    extract_tv_niveis(tv_doc)
    extract_tv_classes(tv_doc)
    extract_tv_protecao(tv_doc)


if __name__ == "__main__":
    run()
